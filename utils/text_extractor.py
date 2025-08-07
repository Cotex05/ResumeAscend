import PyPDF2
import docx
import io
import streamlit as st

def extract_text_from_file(uploaded_file):
    """
    Extract text from uploaded PDF or DOCX file
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    
    file_extension = uploaded_file.name.lower().split('.')[-1]
    
    try:
        if file_extension == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            return extract_text_from_docx(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        st.error(f"Error extracting text from {file_extension.upper()} file: {str(e)}")
        return ""

def extract_text_from_pdf(uploaded_file):
    """
    Extract text from PDF file using PyPDF2
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    
    text_content = ""
    
    try:
        # Read the uploaded file into a BytesIO object
        pdf_bytes = uploaded_file.read()
        pdf_file = io.BytesIO(pdf_bytes)
        
        # Create PDF reader object
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Extract text from all pages
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text_content += page.extract_text() + "\n"
        
        # Clean up the extracted text
        text_content = clean_extracted_text(text_content)
        
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    return text_content

def extract_text_from_docx(uploaded_file):
    """
    Extract text from DOCX file using python-docx
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        str: Extracted text content
    """
    
    text_content = ""
    
    try:
        # Read the uploaded file into a BytesIO object
        docx_bytes = uploaded_file.read()
        docx_file = io.BytesIO(docx_bytes)
        
        # Create Document object
        doc = docx.Document(docx_file)
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"
        
        # Clean up the extracted text
        text_content = clean_extracted_text(text_content)
        
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    return text_content

def clean_extracted_text(text):
    """
    Clean and normalize extracted text
    
    Args:
        text (str): Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    
    if not text:
        return ""
    
    # Remove excessive whitespace and newlines
    import re
    
    # Replace multiple spaces with single space
    text = re.sub(r'\s+', ' ', text)
    
    # Replace multiple newlines with double newline
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

def validate_file_content(text):
    """
    Validate that the extracted text contains meaningful content
    
    Args:
        text (str): Extracted text content
        
    Returns:
        bool: True if content is valid, False otherwise
    """
    
    if not text or len(text.strip()) < 50:
        return False
    
    # Check for common resume indicators
    resume_indicators = [
        'experience', 'education', 'skills', 'work', 'employment',
        'university', 'college', 'degree', 'email', 'phone', 'address'
    ]
    
    text_lower = text.lower()
    indicator_count = sum(1 for indicator in resume_indicators if indicator in text_lower)
    
    # Require at least 2 resume indicators
    return indicator_count >= 2
