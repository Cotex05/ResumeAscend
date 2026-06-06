import io
import docx

class DocumentExporter:
    @staticmethod
    def export_to_docx(markdown_text: str) -> io.BytesIO:
        """Converts simple markdown text to a DOCX file stream."""
        doc = docx.Document()
        lines = markdown_text.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            if line_stripped.startswith('# '):
                doc.add_heading(line_stripped[2:], level=1)
            elif line_stripped.startswith('## '):
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith('### '):
                doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                doc.add_paragraph(line_stripped[2:], style='List Bullet')
            else:
                # Remove bold asterisks for plain text paragraph
                clean_line = line_stripped.replace('**', '').replace('__', '')
                doc.add_paragraph(clean_line)
                
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    @staticmethod
    def _sanitize_for_pdf(text: str) -> str:
        replacements = {
            '\u2013': '-',   # en-dash
            '\u2014': '--',  # em-dash
            '\u2018': "'",   # left single quote
            '\u2019': "'",   # right single quote
            '\u201c': '"',   # left double quote
            '\u201d': '"',   # right double quote
            '\u2022': '*',   # bullet
            '\u2026': '...', # ellipsis
            '\u00a0': ' ',   # non-breaking space
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        # Remove remaining characters that don't fit in windows-1252
        return text.encode('windows-1252', 'ignore').decode('windows-1252')

    @staticmethod
    def export_json_to_pdf(resume_json: dict) -> bytes:
        """Converts structured JSON resume to a PDF byte string using ReportLab."""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        import io

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )

        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'NameTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=5,
            textColor=colors.HexColor('#1a1a1a')
        )
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=20,
            textColor=colors.HexColor('#666666')
        )
        heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=10,
            textColor=colors.HexColor('#2c3e50'),
            borderPadding=(0, 0, 2, 0),
            borderColor=colors.HexColor('#e0e0e0'),
            borderWidth=1
        )
        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=8
        )
        bullet_style = ParagraphStyle(
            'BulletText',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=4
        )

        story = []

        # Helper to safely get and sanitize text
        def safe_text(val):
            if not val: return ""
            return DocumentExporter._sanitize_for_pdf(str(val))

        # Name & Title
        name = safe_text(resume_json.get('name', 'Name Not Provided'))
        title = safe_text(resume_json.get('title', ''))
        story.append(Paragraph(name, title_style))
        if title:
            story.append(Paragraph(title, subtitle_style))
            
        # Summary
        summary = safe_text(resume_json.get('summary', ''))
        if summary:
            story.append(Paragraph('Professional Summary', heading_style))
            story.append(Paragraph(summary, body_style))
            
        # Experience
        experience = resume_json.get('experience', [])
        if experience:
            story.append(Paragraph('Experience', heading_style))
            for job in experience:
                # We assume the AI returns properly formatted strings, possibly with basic markdown or raw text.
                # ReportLab Paragraph supports basic HTML tags like <b> or <i>
                safe_job = safe_text(job).replace('\n', '<br/>')
                story.append(Paragraph(safe_job, body_style))
                story.append(Spacer(1, 5))
                
        # Education
        education = resume_json.get('education', [])
        if education:
            story.append(Paragraph('Education', heading_style))
            for edu in education:
                safe_edu = safe_text(edu).replace('\n', '<br/>')
                story.append(Paragraph(f"• {safe_edu}", bullet_style))
                
        # Skills
        skills = resume_json.get('skills', [])
        if skills:
            story.append(Paragraph('Skills', heading_style))
            skills_text = ", ".join([safe_text(s) for s in skills])
            story.append(Paragraph(skills_text, body_style))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes
